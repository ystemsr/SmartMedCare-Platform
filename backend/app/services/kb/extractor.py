"""Extract plain text from supported KB source files.

Supported formats: md, pdf, csv, txt, docx.
"""

from __future__ import annotations

import csv
import io
import logging
from typing import Callable, Dict

logger = logging.getLogger(__name__)


SUPPORTED_EXTENSIONS = {"md", "markdown", "pdf", "csv", "txt", "docx"}


def detect_extension(file_name: str) -> str:
    """Return the lowercase extension without the leading dot."""
    dot = file_name.rfind(".")
    if dot < 0 or dot == len(file_name) - 1:
        return ""
    return file_name[dot + 1 :].lower()


def _extract_txt(data: bytes) -> str:
    for enc in ("utf-8", "gbk", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_md(data: bytes) -> str:
    # Markdown is plain text; the chunker eats the formatting well enough
    # without an AST pass. Strip common boilerplate to reduce noise.
    text = _extract_txt(data)
    return text


def _extract_csv(data: bytes) -> str:
    """Flatten CSV rows into pipe-delimited lines so chunks retain context."""
    text = _extract_txt(data)
    reader = csv.reader(io.StringIO(text))
    lines = []
    header: list[str] | None = None
    for row in reader:
        if not row:
            continue
        if header is None:
            header = row
            lines.append(" | ".join(header))
            continue
        # Pair each value with its header so chunked rows stay interpretable.
        pairs = []
        for i, cell in enumerate(row):
            key = header[i] if i < len(header) else f"col{i}"
            pairs.append(f"{key}: {cell}")
        lines.append("; ".join(pairs))
    return "\n".join(lines)


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception as e:  # noqa: BLE001
            logger.warning("pdf page %s extract failed: %s", i, e)
            text = ""
        if text.strip():
            parts.append(text)
    return "\n\n".join(parts)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    # Tables as tab-separated rows.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts)


_EXTRACTORS: Dict[str, Callable[[bytes], str]] = {
    "txt": _extract_txt,
    "md": _extract_md,
    "markdown": _extract_md,
    "csv": _extract_csv,
    "pdf": _extract_pdf,
    "docx": _extract_docx,
}


class UnsupportedFileType(ValueError):
    pass


def extract_text(file_name: str, data: bytes) -> str:
    """Extract plain text from a supported file. Raises on unsupported types."""
    ext = detect_extension(file_name)
    fn = _EXTRACTORS.get(ext)
    if fn is None:
        raise UnsupportedFileType(
            f"Unsupported file type: .{ext or '?'} — allowed: "
            f"{', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    return fn(data)
